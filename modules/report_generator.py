from pathlib import Path
from datetime import datetime
import pandas as pd
from fpdf import FPDF
from docx import Document
import io

REPORT_DIR = Path('reports')

def generate_text_report(filename, metrics, findings_df, risk_level, risk_score, suggestions, migrated_code=""):
    REPORT_DIR.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    report_path = REPORT_DIR / f"migration_report_{timestamp}.txt"
    
    issues_text = findings_df.to_string(index=False) if not findings_df.empty else 'No issues found.'
    content = f"""
AI-Based Legacy Code Migration & Analysis System
================================================
File Name: {filename}
Risk Level: {risk_level}
Risk Score: {risk_score}%
Code Metrics:
- Total Lines: {metrics.get('total_lines')}
- Functions: {metrics.get('functions')}
- Classes: {metrics.get('classes')}
Detected Issues:
{issues_text}
Migration Suggestions:
{suggestions}
Generated At: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
    report_path.write_text(content, encoding='utf-8')
    return report_path, content

def _strip_nonprintable(s: str) -> str:
    # Remove control chars that can break PDF layout/rendering.
    return ''.join(ch for ch in s if ch == '\n' or ch == '\t' or (32 <= ord(ch) <= 126))


def clean_text(value, max_len: int = 20000) -> str:
    """Convert arbitrary Python values into PDF-safe printable text.

    - Handles: list/tuple/set/dict, None, numbers, objects.
    - Prevents crashes by ensuring the result is always a str.
    - Removes emojis/unicode symbols by keeping ASCII printable characters only.
    """
    try:
        if value is None:
            return ""

        # Convert common complex types deterministically
        if isinstance(value, dict):
            lines = []
            for k, v in value.items():
                lines.append(f"{k}: {v}")
            value = "\n".join(lines)

        elif isinstance(value, (list, tuple, set)):
            # Join each element on newline; keep nesting readable.
            value = "\n".join(clean_text(v, max_len=max_len) for v in value)

        elif isinstance(value, bytes):
            value = value.decode('utf-8', errors='replace')

        # Fallback: coerce to string
        text = str(value)

        # Normalize newlines
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # Remove emojis/unicode symbols by mapping non-ascii to '?'
        # This avoids font/encoding issues in default FPDF.
        text = text.encode('ascii', errors='replace').decode('ascii', errors='replace')

        # Strip non-printable/control characters but keep newlines/tabs
        text = _strip_nonprintable(text)

        if max_len and len(text) > max_len:
            text = text[:max_len] + "\n...[truncated]"

        return text

    except Exception:
        # Last resort: never crash the PDF generator.
        try:
            return "[unprintable value]"
        except Exception:
            return ""


def _pdf_write_title(pdf: FPDF, title: str) -> None:
    pdf.set_font("helvetica", 'B', 16)
    pdf.cell(0, 10, txt=title, ln=True, align='C')
    pdf.ln(2)
    pdf.set_draw_color(0, 120, 200)
    pdf.set_line_width(0.4)
    # Horizontal line (full width)
    x = pdf.l_margin
    y = pdf.get_y()
    page_w = pdf.w - pdf.l_margin - pdf.r_margin
    pdf.line(x, y, x + page_w, y)
    pdf.ln(6)


def _pdf_write_heading(pdf: FPDF, heading: str) -> None:
    pdf.set_font("helvetica", 'B', 13)
    pdf.set_text_color(0, 60, 110)
    pdf.cell(0, 8, txt=heading, ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(1)


def _pdf_write_kv(pdf: FPDF, label: str, value: str) -> None:
    # Two-column key/value row (fixed label width)
    label = clean_text(label)
    value = clean_text(value)
    label_w = 55
    pdf.set_font("helvetica", '', 10)
    pdf.cell(label_w, 6, txt=label, ln=0)
    # Multi_cell for value to wrap
    cur_x = pdf.get_x()
    cur_y = pdf.get_y()
    pdf.set_xy(cur_x + label_w, cur_y)
    pdf.set_font("helvetica", '', 10)
    pdf.multi_cell(0, 6, txt=value)


def _pdf_write_block(pdf: FPDF, text: str, font_size: int = 10, line_height: int = 5) -> None:
    # Multi-cell with cleaned text and automatic page breaks.
    pdf.set_font("helvetica", '', font_size)
    cleaned = clean_text(text)
    if not cleaned:
        pdf.set_font("helvetica", 'I', font_size)
        pdf.cell(0, line_height, txt="(none)", ln=True)
        return

    # Ensure the block wraps safely.
    for para in cleaned.split("\n"):
        pdf.multi_cell(0, line_height, txt=para)


def generate_pdf_report(data):
    """Generate a PDF report (bytes) safe for Streamlit downloads.

    Defensive coding: never allow Unicode/emoji/list/dict/None values
    to crash FPDF. All content is cleaned via `clean_text()`.
    """
    try:
        filename = clean_text(data.get('filename', 'unknown'))
        risk_level = clean_text(data.get('risk_level', 'N/A'))
        risk_score = data.get('risk_score', '')
        metrics = data.get('metrics', {}) or {}

        # Some pipelines use `suggestions_text`, some older ones may use `suggestions`.
        suggestions = data.get('suggestions_text', None)
        if suggestions is None:
            suggestions = data.get('suggestions', None)

        issues = data.get('findings_df', None)
        # Convert findings_df to lines early to keep PDF fast.
        issues_text = ''
        try:
            if issues is not None and hasattr(issues, 'empty') and not issues.empty:
                # to_string can still include unicode; we'll clean later.
                issues_text = issues.to_string(index=False)
            else:
                issues_text = 'No issues found.'
        except Exception:
            issues_text = 'No issues found.'

        pdf = FPDF(format='A4')
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        _pdf_write_title(pdf, "Code Migration Analysis Report")

        # Header metadata
        pdf.set_font("helvetica", '', 11)
        pdf.cell(0, 7, txt=f"Filename: {filename}", ln=True)
        pdf.cell(0, 7, txt=f"Risk Level: {risk_level} ({clean_text(risk_score)}%)", ln=True)
        pdf.ln(3)

        # Metrics table-like section
        _pdf_write_heading(pdf, "Metrics")
        for k in [
            'total_lines', 'non_empty_lines', 'functions', 'classes', 'imports', 'comments'
        ]:
            if k in metrics:
                _pdf_write_kv(pdf, k.replace('_', ' ').title(), metrics.get(k, ''))
        # If metrics dict is empty, still print something.
        if not metrics:
            _pdf_write_block(pdf, "(no metrics)")

        pdf.ln(2)

        # Detected issues
        _pdf_write_heading(pdf, "Detected Issues")
        _pdf_write_block(pdf, issues_text, font_size=9, line_height=4)

        pdf.ln(2)

        # Suggestions section
        _pdf_write_heading(pdf, "Migration Suggestions")
        _pdf_write_block(pdf, suggestions, font_size=10, line_height=5)

        # Optional diff/roadmap
        roadmap = data.get('roadmap_text')
        if roadmap:
            pdf.ln(2)
            _pdf_write_heading(pdf, "Migration Roadmap")
            _pdf_write_block(pdf, roadmap, font_size=10, line_height=5)

        # Return PDF bytes for Streamlit download_button
        # Using dest='S' returns a byte string.
        out = pdf.output(dest='S')
        if isinstance(out, str):
            out = out.encode('latin-1', errors='ignore')
        return out

    except Exception as e:
        # Never crash the Streamlit app: return a minimal PDF with the error message.
        pdf = FPDF(format='A4')
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        _pdf_write_title(pdf, "Code Migration Analysis Report")
        _pdf_write_heading(pdf, "PDF Generation Error")
        pdf.set_font("helvetica", '', 10)
        pdf.multi_cell(0, 5, txt=clean_text(str(e)))
        out = pdf.output(dest='S')
        if isinstance(out, str):
            out = out.encode('latin-1', errors='ignore')
        return out


def generate_docx_report(data):
    doc = Document()
    doc.add_heading('Code Migration Analysis Report', 0)
    
    p = doc.add_paragraph()
    p.add_run(f"Filename: ").bold = True
    p.add_run(data.get('filename', 'unknown'))
    
    doc.add_heading('Metrics', level=1)
    for k, v in data.get('metrics', {}).items():
        doc.add_paragraph(f"{k.replace('_', ' ').title()}: {v}", style='List Bullet')
        
    doc.add_heading('Suggestions', level=1)
    doc.add_paragraph(data.get('suggestions_text', data.get('suggestions', 'No suggestions available.')))
    
    target = io.BytesIO()
    doc.save(target)
    return target.getvalue()

def generate_html_report(data):
    """Generates a professional HTML representation of the analysis."""
    findings = data.get('findings_df')
    issues_html = findings.to_html(classes='table table-striped') if findings is not None and not findings.empty else '<p>No issues found.</p>'
    html_content = f"""
    <html>
    <head>
        <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css">
    </head>
    <body class="container py-5">
        <h1 class="mb-4">Code Migration Report: {data.get('filename', 'unknown')}</h1>
        <div class="alert alert-info">Risk Level: {data.get('risk_level', 'N/A')} | Score: {data.get('risk_score', 'N/A')}%</div>
        <h2>Metrics</h2>
        <ul>
            {" ".join([f"<li>{k.replace('_', ' ').title()}: {v}</li>" for k, v in data.get('metrics', {}).items()])}
        </ul>
        <h2>Detected Issues</h2>
        {issues_html}
    </body>
    </html>
    """
    return html_content

def export_findings_csv(findings_df: pd.DataFrame):
    REPORT_DIR.mkdir(exist_ok=True)
    csv_path = REPORT_DIR / f"findings_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    findings_df.to_csv(csv_path, index=False)
    return csv_path
