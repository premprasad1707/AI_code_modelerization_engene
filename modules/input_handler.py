"""
modules/input_handler.py
Handles multi-format inputs: ZIP archives, PDF files, Word DOCX documents.
"""

import io
import re
import zipfile
import logging
from typing import Dict

logger = logging.getLogger(__name__)

SUPPORTED_CODE_EXTENSIONS = {
    ".py", ".java", ".js", ".ts", ".php", ".c", ".cpp",
    ".sql", ".sh", ".rb", ".go", ".rs", ".kt", ".cs",
    ".html", ".css", ".xml", ".yaml", ".yml", ".json", ".txt",
}


# ── ZIP handling ───────────────────────────────────────────────────────────────

def handle_zip_upload(zip_file) -> Dict[str, str]:
    """
    Extract supported code files from a ZIP upload.

    Args:
        zip_file: Streamlit UploadedFile or file-like object

    Returns:
        dict of {relative_path: code_text}
    """
    files = {}
    try:
        # Ensure we are at the start of the file if it was read previously
        if hasattr(zip_file, "seek"):
            zip_file.seek(0)
        with zipfile.ZipFile(zip_file) as zf:
            for name in zf.namelist():
                # Skip directories, hidden files, __pycache__, .git
                if name.endswith("/") or "/." in name or "__pycache__" in name or ".git/" in name:
                    continue
                ext = "." + name.rsplit(".", 1)[-1].lower() if "." in name else ""
                if ext in SUPPORTED_CODE_EXTENSIONS:
                    try:
                        content = zf.read(name).decode("utf-8", errors="ignore")
                        if content.strip():
                            files[name] = content
                    except Exception as e:
                        logger.warning(f"Could not read {name} from ZIP: {e}")
    except zipfile.BadZipFile:
        logger.error("Invalid ZIP file uploaded.")
    logger.info(f"ZIP extraction: {len(files)} code files found.")
    return files


# ── PDF handling ───────────────────────────────────────────────────────────────

def extract_code_from_pdf(pdf_file) -> str:
    """
    Extract text (including code blocks) from a PDF.
    Uses pdfplumber if available, falls back to PyPDF2.

    Returns extracted text as a string.
    """
    content = ""
    raw_bytes = pdf_file.read()

    # Try pdfplumber first (better layout preservation)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
            content = "\n".join(pages_text)
        logger.info(f"PDF extracted via pdfplumber: {len(content)} chars")
        return content
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}; trying PyPDF2")

    # Fallback to PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(raw_bytes))
        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        content = "\n".join(pages_text)
        logger.info(f"PDF extracted via PyPDF2: {len(content)} chars")
        return content
    except ImportError:
        logger.error("Neither pdfplumber nor PyPDF2 is installed.")
        return "# PDF extraction failed: install pdfplumber or PyPDF2\n"
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return f"# PDF extraction error: {e}\n"


# ── DOCX handling ──────────────────────────────────────────────────────────────

def extract_code_from_docx(docx_file) -> str:
    """
    Extract text content (paragraphs + table cells) from a Word .docx file.
    Preserves paragraph spacing for code-like blocks.

    Returns extracted text as a string.
    """
    try:
        from docx import Document
        raw_bytes = docx_file.read()
        doc = Document(io.BytesIO(raw_bytes))

        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)

        # Also extract table cell text (code often appears in tables in docs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())

        content = "\n".join(lines)
        logger.info(f"DOCX extracted: {len(content)} chars, {len(lines)} lines")
        return content
    except ImportError:
        logger.error("python-docx is not installed.")
        return "# DOCX extraction failed: install python-docx\n"
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return f"# DOCX extraction error: {e}\n"
